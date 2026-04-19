"""
Ingestor module — handles all document loading, outline extraction, and chunking.
Supports: PDF (.pdf), DOCX (.docx / .doc)

Pipeline per file:
  1. Load raw text page-by-page (PDF) or paragraph-by-paragraph (DOCX)
  2. Extract structural outline (headings / bookmarks)
  3. Split text into overlapping chunks using RecursiveCharacterTextSplitter
  4. Attach nearest heading + source filename + page number to every chunk

Output schema per chunk:
  {
      "text":    str,   # chunk content
      "source":  str,   # original filename (not full path)
      "page":    int,   # 1-indexed page number (DOCX uses logical page estimate)
      "heading": str,   # nearest section heading, "" if none found
  }
"""

import re
from pathlib import Path
from typing import Any

import pypdf

# LangChain 0.2+ moved splitters to langchain_text_splitters;
# fall back to the old path for older installs.
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter

# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

SUPPORTED = {".pdf", ".docx", ".doc"}


def supported_formats() -> list[str]:
    """Return list of accepted file extensions."""
    return sorted(SUPPORTED)


def extract_outline(file_path: str) -> list[str]:
    """
    Extract the heading / bookmark structure from a document.
    Returns a flat, ordered list of heading strings.

    PDF  — walks the outline/bookmarks tree returned by pypdf.
    DOCX — iterates paragraphs whose style starts with 'Heading'.
    """
    path = Path(file_path)
    _validate_format(path)

    if path.suffix.lower() == ".pdf":
        return _pdf_outline(path)
    else:
        return _docx_outline(path)


def load_and_chunk(
    file_path: str,
    chunk_size: int = 500,
    overlap: int = 100,
) -> tuple[list[dict], list[str]]:
    """
    Load a document, extract its outline, and split into chunks.

    Args:
        file_path:  Absolute or relative path to the file.
        chunk_size: Target character count per chunk (default 500).
        overlap:    Character overlap between consecutive chunks (default 100).

    Returns:
        chunks:  list of chunk dicts — see module docstring for schema.
        outline: list of heading strings in document order.

    Raises:
        ValueError:       Unsupported file format.
        FileNotFoundError: File does not exist.
    """
    path = Path(file_path)
    _validate_format(path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    outline = extract_outline(file_path)

    # Load raw pages: list of {"text": str, "page": int}
    if path.suffix.lower() == ".pdf":
        pages = _load_pdf_pages(path)
    else:
        pages = _load_docx_pages(path)

    # Build a heading-position map for attaching headings to chunks
    # heading_positions: list of (char_offset, heading_str) across full text
    heading_positions = _build_heading_positions(pages, outline, path)

    # Flatten all pages into a single text with page-boundary markers so we
    # can recover page numbers after splitting.
    flat_text, offset_to_page = _flatten_pages(pages)

    # Split
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    raw_chunks = splitter.split_text(flat_text)

    # Rebuild chunks with metadata
    chunks: list[dict] = []
    cursor = 0  # track position in flat_text to map back to page/heading

    source_name = path.name

    for raw in raw_chunks:
        # Find this chunk's start position in flat_text
        start = flat_text.find(raw, cursor)
        if start == -1:
            start = cursor  # fallback; shouldn't happen

        page_num = _page_at_offset(start, offset_to_page)
        heading  = _heading_at_offset(start, heading_positions)

        chunks.append({
            "text":    raw.strip(),
            "source":  source_name,
            "page":    page_num,
            "heading": heading,
        })

        cursor = max(cursor, start + max(1, len(raw) - overlap))

    return chunks, outline


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

def _load_pdf_pages(path: Path) -> list[dict]:
    """Return list of {"text": str, "page": int} for every PDF page."""
    pages = []
    reader = pypdf.PdfReader(str(path))
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append({"text": text, "page": i})
    return pages


def _pdf_outline(path: Path) -> list[str]:
    """
    Walk pypdf's outline tree recursively and return flat list of titles.
    pypdf returns a mixed list of Destination objects and nested lists.
    """
    reader = pypdf.PdfReader(str(path))
    outline = reader.outline  # may be [] if no bookmarks

    headings: list[str] = []
    _walk_pdf_outline(outline, headings)
    return headings


def _walk_pdf_outline(node: Any, result: list[str]) -> None:
    """Recursively walk a pypdf outline node."""
    if isinstance(node, list):
        for item in node:
            _walk_pdf_outline(item, result)
    elif hasattr(node, "title"):
        title = node.title.strip()
        if title:
            result.append(title)


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _load_docx_pages(path: Path) -> list[dict]:
    """
    docx2txt doesn't expose page breaks cleanly.
    We use python-docx (bundled as part of docx2txt's dependency) to iterate
    paragraphs and estimate page numbers via page-break runs.

    Falls back to docx2txt for body text if python-docx is unavailable.
    """
    try:
        import docx  # python-docx
        return _load_docx_via_python_docx(path)
    except ImportError:
        # Fallback: treat entire doc as one page
        import docx2txt  # lazy import — only needed in fallback path
        text = docx2txt.process(str(path)) or ""
        return [{"text": text, "page": 1}]


def _load_docx_via_python_docx(path: Path) -> list[dict]:
    """Use python-docx to iterate paragraphs and detect page breaks."""
    import docx
    from docx.oxml.ns import qn

    doc    = docx.Document(str(path))
    pages  = []
    current_page = 1
    current_text: list[str] = []

    for para in doc.paragraphs:
        # Detect explicit page break inside paragraph runs
        for run in para.runs:
            if run._element.xml.find("w:br") != -1:
                br_elements = run._element.findall(
                    f".//{{{run._element.nsmap.get('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')}}}br"
                )
                for br in br_elements:
                    br_type = br.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type",
                        ""
                    )
                    if br_type == "page":
                        pages.append({"text": "\n".join(current_text), "page": current_page})
                        current_page += 1
                        current_text = []

        text = para.text.strip()
        if text:
            current_text.append(text)

    if current_text:
        pages.append({"text": "\n".join(current_text), "page": current_page})

    if not pages:
        pages = [{"text": "", "page": 1}]

    return pages


def _docx_outline(path: Path) -> list[str]:
    """
    Extract headings from DOCX by inspecting paragraph styles.
    Paragraphs with style name starting with 'Heading' are included.
    """
    try:
        import docx
        doc = docx.Document(str(path))
        headings = []
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading") and para.text.strip():
                headings.append(para.text.strip())
        return headings
    except ImportError:
        # python-docx not available — return empty outline
        return []


# ---------------------------------------------------------------------------
# Shared utilities
# ---------------------------------------------------------------------------

def _validate_format(path: Path) -> None:
    """Raise ValueError if file extension is not supported."""
    if path.suffix.lower() not in SUPPORTED:
        raise ValueError(
            f"Unsupported file format '{path.suffix}'. "
            f"Supported: {', '.join(sorted(SUPPORTED))}"
        )


def _flatten_pages(pages: list[dict]) -> tuple[str, list[tuple[int, int]]]:
    """
    Concatenate all page texts into one string.

    Returns:
        flat_text:       Full document text as a single string.
        offset_to_page:  List of (start_offset, page_number) tuples,
                         sorted by offset. Use binary search or linear scan
                         to map an offset → page number.
    """
    parts: list[str] = []
    offset_to_page: list[tuple[int, int]] = []
    cursor = 0

    for p in pages:
        text = p["text"]
        offset_to_page.append((cursor, p["page"]))
        parts.append(text)
        cursor += len(text) + 1  # +1 for the joining newline

    flat_text = "\n".join(parts)
    return flat_text, offset_to_page


def _page_at_offset(offset: int, offset_to_page: list[tuple[int, int]]) -> int:
    """Return the page number for a character offset in the flat text."""
    page = 1
    for start, pnum in offset_to_page:
        if offset >= start:
            page = pnum
        else:
            break
    return page


def _build_heading_positions(
    pages: list[dict],
    outline: list[str],
    path: Path,
) -> list[tuple[int, str]]:
    """
    Build a sorted list of (char_offset_in_flat_text, heading_str).

    For PDFs we scan the flat text for each heading string.
    For DOCX we do the same (since python-docx already gave us headings).
    """
    _, offset_to_page = _flatten_pages(pages)
    flat_text, _ = _flatten_pages(pages)

    positions: list[tuple[int, str]] = []
    search_start = 0

    for heading in outline:
        idx = flat_text.find(heading, search_start)
        if idx != -1:
            positions.append((idx, heading))
            search_start = idx + len(heading)

    positions.sort(key=lambda x: x[0])
    return positions


def _heading_at_offset(
    offset: int,
    heading_positions: list[tuple[int, str]],
) -> str:
    """
    Return the heading that most recently precedes `offset`.
    Returns "" if no heading found before this position.
    """
    result = ""
    for pos, heading in heading_positions:
        if pos <= offset:
            result = heading
        else:
            break
    return result
