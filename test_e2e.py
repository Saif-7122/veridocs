"""
End-to-end test: create sample DOCX files, upload, chat, insights, compare, report, delete.
"""
import requests, os, time
from docx import Document

BASE = "http://127.0.0.1:8000"
TEST_DIR = os.path.join(os.path.dirname(__file__), "test_samples")
os.makedirs(TEST_DIR, exist_ok=True)

# ── Step 0: Create two sample DOCX files ─────────────────────────────────
print("=" * 60)
print("STEP 0: Creating sample DOCX files...")

doc1 = Document()
doc1.add_heading("Artificial Intelligence in Healthcare", level=1)
doc1.add_paragraph(
    "Artificial intelligence (AI) is transforming the healthcare industry. "
    "Machine learning models can now detect diseases from medical images with "
    "accuracy that rivals human radiologists. Natural language processing is "
    "being used to extract insights from electronic health records. AI-powered "
    "chatbots are providing preliminary diagnoses and triage services. Deep "
    "learning algorithms are accelerating drug discovery by predicting molecular "
    "interactions. The integration of AI in healthcare promises to improve patient "
    "outcomes, reduce costs, and make medical services more accessible globally."
)
doc1.add_heading("Key Applications", level=2)
doc1.add_paragraph("1. Medical imaging and diagnostics")
doc1.add_paragraph("2. Drug discovery and development")
doc1.add_paragraph("3. Personalized treatment plans")
doc1.add_paragraph("4. Electronic health record analysis")
doc1.add_paragraph("5. Remote patient monitoring")
path1 = os.path.join(TEST_DIR, "ai_healthcare.docx")
doc1.save(path1)
print(f"  Created: {path1}")

doc2 = Document()
doc2.add_heading("Blockchain Technology in Finance", level=1)
doc2.add_paragraph(
    "Blockchain technology is revolutionizing the financial sector. Decentralized "
    "ledgers enable transparent and immutable transaction records. Smart contracts "
    "automate complex financial agreements without intermediaries. Cryptocurrency "
    "exchanges are providing alternative investment opportunities. Central banks "
    "around the world are exploring Central Bank Digital Currencies (CBDCs). "
    "The technology reduces fraud, lowers transaction costs, and increases the "
    "speed of cross-border payments. However, regulatory challenges and energy "
    "consumption concerns remain significant hurdles for widespread adoption."
)
doc2.add_heading("Key Applications", level=2)
doc2.add_paragraph("1. Decentralized finance (DeFi)")
doc2.add_paragraph("2. Cross-border payments")
doc2.add_paragraph("3. Supply chain tracking")
doc2.add_paragraph("4. Digital identity verification")
doc2.add_paragraph("5. Tokenization of assets")
path2 = os.path.join(TEST_DIR, "blockchain_finance.docx")
doc2.save(path2)
print(f"  Created: {path2}")

# ── Step 1: Upload ───────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("STEP 1: Uploading documents...")
with open(path1, "rb") as f1, open(path2, "rb") as f2:
    resp = requests.post(
        f"{BASE}/api/v1/upload",
        files=[
            ("files", ("ai_healthcare.docx", f1, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
            ("files", ("blockchain_finance.docx", f2, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
        ]
    )

print(f"  Status: {resp.status_code}")
if resp.status_code != 200:
    print(f"  ERROR: {resp.text}")
    exit(1)

data = resp.json()
session_id = data["session_id"]
print(f"  Session ID: {session_id}")
print(f"  Files: {data['files']}")
print(f"  Outlines: {list(data['outlines'].keys())}")
print(f"  Status: {data['status']}")

# ── Step 2: Chat ─────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("STEP 2: Testing chat (asking a question)...")
resp = requests.post(
    f"{BASE}/api/v1/chat",
    json={"session_id": session_id, "query": "What are the key applications of AI in healthcare?"}
)
print(f"  Status: {resp.status_code}")
if resp.status_code == 200:
    chat_data = resp.json()
    answer = chat_data.get("answer", "")
    print(f"  Answer (first 300 chars): {answer[:300]}...")
    print(f"  Citations: {chat_data.get('citations', [])}")
else:
    print(f"  ERROR: {resp.text}")

# ── Step 3: Insights ─────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("STEP 3: Extracting insights/themes...")
resp = requests.get(f"{BASE}/api/v1/insights/{session_id}")
print(f"  Status: {resp.status_code}")
if resp.status_code == 200:
    insights_data = resp.json()
    for fname, themes in insights_data.get("themes", {}).items():
        print(f"  {fname}: {str(themes)[:200]}")
else:
    print(f"  ERROR: {resp.text}")

# ── Step 4: Compare ──────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("STEP 4: Comparing documents...")
resp = requests.post(
    f"{BASE}/api/v1/compare",
    json={"session_id": session_id}
)
print(f"  Status: {resp.status_code}")
if resp.status_code == 200:
    compare_data = resp.json()
    print(f"  Comparison result (first 500 chars): {str(compare_data)[:500]}")
else:
    print(f"  ERROR: {resp.text}")

# ── Step 5: Report ───────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("STEP 5: Downloading report...")
resp = requests.get(f"{BASE}/api/v1/report/{session_id}")
print(f"  Status: {resp.status_code}")
if resp.status_code == 200:
    report = resp.text
    print(f"  Report length: {len(report)} chars")
    print(f"  Report preview:\n{report[:500]}")
else:
    print(f"  ERROR: {resp.text}")

# ── Step 6: Delete session ───────────────────────────────────────────────
print("\n" + "=" * 60)
print("STEP 6: Deleting session...")
resp = requests.delete(f"{BASE}/api/v1/session/{session_id}")
print(f"  Status: {resp.status_code}")
print(f"  Response: {resp.json()}")

# ── Summary ──────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("E2E TEST COMPLETE!")
print("=" * 60)
